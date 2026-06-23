# -*- coding: utf-8 -*-
"""내용증명(미수금 독촉) 생성기 — 관할별 명판+인감, 변수 채움. 전부 검정·단일 글꼴·균일 줄간격."""
import os, datetime, re, difflib
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_ALIGN_VERTICAL as VAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F="맑은 고딕"; LS=1.23; BASE=9.5; SHADE="F2F2F2"; BORD="808080"
KWON={
 "서울": dict(corp="충해전기 주식회사", tel="010-5228-2922",
   addr="서울특별시 구로구 구로중앙로198, B-13, 103", stamp="stamp_guro.png",
   bank="신한은행", acct="100-001-944331", holder="충해전기(주)"),
 "화성": dict(corp="충해전기 주식회사 화성영업소", tel="010-5228-2922",
   addr="경기도 화성시 팔탄면 푸른들판로 642-5, G-103~104, 202", stamp="stamp_hwaseong.png",
   bank="신한은행", acct="140-013-471920", holder="충해전기(주) 화성영업소"),
}
SENDER_NAME="김재성"

def won_hangul(n):
    n=int(round(n))
    if n==0: return "영원"
    units=["","만","억","조"]; digit=["","일","이","삼","사","오","육","칠","팔","구"]; pos=["","십","백","천"]
    grp=[]
    while n>0: grp.append(n%10000); n//=10000
    parts=[]
    for i in range(len(grp)-1,-1,-1):
        g=grp[i]
        if g==0: continue
        gs=""
        for p in range(3,-1,-1):
            d=(g//(10**p))%10
            if d==0: continue
            gs += (pos[p] if (d==1 and p>0) else digit[d]+pos[p])
        parts.append(gs+units[i])
    return "".join(parts)+"원"

def _norm_name(s):
    s=re.sub(r"\([^)]*\)","",str(s))
    s=re.sub(r"주식회사|유한회사|㈜|\(주\)|\(유\)|\(자\)","",s)
    return re.sub(r"[^가-힣0-9]","",s)

def load_book(repo_dir):
    """업체명단(_업체명단.xlsx) → {사업자번호: [ {name,tel,rep,addr}, ... ]}. 컬럼: 업체명·전화·팩스·대표자·주소·사업자번호."""
    from openpyxl import load_workbook
    p=os.path.join(repo_dir,"_업체명단.xlsx")
    book={}
    if not os.path.exists(p): return book
    ws=load_workbook(p, read_only=True, data_only=True).active
    for r in ws.iter_rows(min_row=2, values_only=True):
        if not r or r[0] is None: continue
        biz=str(r[5] or "").strip() if len(r)>5 else ""
        if not re.match(r"\d{3}-\d{2}-\d{5}", biz): continue
        book.setdefault(biz,[]).append({"name":str(r[0]).strip(),"tel":str(r[1] or "").strip(),
            "rep":str(r[3] or "").strip(),"addr":str(r[4] or "").strip()})
    return book

def match_company(book, name, biz):
    """사업자번호로 찾고, 같은 번호가 여러 곳이면 이름으로 구분. 없으면 None."""
    c=book.get(biz) or []
    if not c: return None
    if len(c)==1: return c[0]
    dn=_norm_name(name)
    for o in c:
        if _norm_name(o["name"])==dn: return o
    return max(c, key=lambda o:difflib.SequenceMatcher(None,dn,_norm_name(o["name"])).ratio())

def make_docx(data, repo_dir, out_docx, font=None, date_before=12, item_after=3, spread=0):
    FF=font or F
    reg=data["관할"]; K=KWON[reg]; amt=int(round(data["미수액"]))
    dt=data.get("날짜") or datetime.date.today()
    if isinstance(dt,str):
        try: dt=datetime.date.fromisoformat(dt)
        except Exception: dt=datetime.date.today()
    datestr=f"{dt.year}년 {dt.month}월 {dt.day}일"
    doc=Document()
    st=doc.styles['Normal']; st.font.name=FF; st.font.size=Pt(BASE); st.element.rPr.rFonts.set(qn('w:eastAsia'),FF)
    pf0=st.paragraph_format; pf0.line_spacing=LS
    sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(1.1); sec.bottom_margin=Cm(0.9); sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
    CW=Cm(21-4.0)
    def kr(run,size=BASE,bold=False):
        run.font.name=FF; run.font.size=Pt(size); run.font.bold=bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'),FF); return run
    def para(text="",align=None,size=BASE,bold=False,before=0,after=3):
        p=doc.add_paragraph()
        if align is not None: p.alignment=align
        pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=LS
        if text!="": kr(p.add_run(text),size,bold)
        return p
    def numitem(num, text, after=None, before=0):
        p=doc.add_paragraph(); pf=p.paragraph_format
        pf.space_after=Pt(item_after if after is None else after); pf.space_before=Pt(before); pf.line_spacing=LS
        kr(p.add_run(num+". "),BASE); kr(p.add_run(text),BASE); return p
    def shade(cell,fill=SHADE):
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
    def tblb(t):
        tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
        for ed in ('top','left','bottom','right','insideH','insideV'):
            e=OxmlElement('w:'+ed); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:space'),'0'); e.set(qn('w:color'),BORD); b.append(e)
        tblPr.append(b)
    def cp(cell,text,size=BASE,bold=False,align=None):
        p=cell.paragraphs[0]
        if align is not None: p.alignment=align
        p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0); p.paragraph_format.line_spacing=LS
        if text!="": kr(p.add_run(text),size,bold)
        cell.vertical_alignment=VAL.CENTER; return p
    def rh(row,cm):
        trPr=row._tr.get_or_add_trPr(); h=OxmlElement('w:trHeight'); h.set(qn('w:val'),str(int(cm*567))); h.set(qn('w:hRule'),'atLeast'); trPr.append(h)
    def hrule(p,sz="6"):
        pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); bb=OxmlElement('w:bottom')
        bb.set(qn('w:val'),'single'); bb.set(qn('w:sz'),sz); bb.set(qn('w:space'),'3'); bb.set(qn('w:color'),"000000"); pb.append(bb); pPr.append(pb)
    def set_fixed(t, widths_cm):
        tbl=t._tbl; tblPr=tbl.tblPr
        lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
        w=OxmlElement('w:tblW'); w.set(qn('w:w'),str(int(sum(widths_cm)*567))); w.set(qn('w:type'),'dxa'); tblPr.append(w)
        grid=tbl.find(qn('w:tblGrid'))
        for gc in list(grid): grid.remove(gc)
        for wcm in widths_cm:
            gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(int(wcm*567))); grid.append(gc)
        for row in t.rows:
            for i,c in enumerate(row.cells): c.width=Cm(widths_cm[i])
    def info_table(rows):
        t=doc.add_table(rows=2,cols=4); tblb(t)
        ws=[Cm(2.2),Cm(5.3),Cm(2.2),Cm(CW.cm-9.7)]
        c=t.row_cells(0)
        for i,w in enumerate(ws): c[i].width=w
        shade(c[0]); shade(c[2])
        cp(c[0],rows[0][0],bold=True,align=AL.CENTER); cp(c[1],rows[0][1]); cp(c[2],rows[0][2],bold=True,align=AL.CENTER); cp(c[3],rows[0][3])
        c1=t.row_cells(1)
        for i,w in enumerate(ws): c1[i].width=w
        shade(c1[0]); cp(c1[0],rows[1][0],bold=True,align=AL.CENTER)
        m=c1[1].merge(c1[2]).merge(c1[3]); cp(m,rows[1][1])
        for r in t.rows: rh(r,0.68)

    para("내  용  증  명  서",AL.CENTER,size=22,bold=True,before=2,after=8)
    para("수신자",size=10.5,bold=True,before=2,after=2)
    info_table([("성  명",data["담당자"],"전화번호",data["수신전화"]),("주  소",data["수신주소"])])
    para("발신자",size=10.5,bold=True,before=5+spread,after=2)
    info_table([("성  명",SENDER_NAME,"전화번호",K["tel"]),("주  소",K["addr"])])
    hrule(para("미수금 지불 촉구의 건",size=10.5,bold=True,before=7+spread,after=4))
    numitem("1","귀하(社)의 무궁한 발전을 기원합니다.",before=11)
    numitem("2",f'본 내용증명은 발신인 {K["corp"]} (이하 “발신인”) 가 수신인 {data["거래처명"]} (이하 “수신인”) 에게 납품한 유공압 제품의 미지급 대금에 관한 것입니다.')
    numitem("3",f'발신인은 수신인의 주문에 따라 유공압 제품을 정상적으로 납품하였으며, 이에 대한 납품대금은 부가세를 포함하여 금{amt:,}원({won_hangul(amt)})입니다.')
    numitem("4","양 당사자는 위 납품대금을 약정 지급일까지 지급하기로 약정하였으나, 수신인은 약정 지급기일이 경과한 본 서면 작성일 현재까지도 위 대금을 지급하지 아니하고 있습니다.")
    numitem("5",f'이에 발신인은 수신인에게 본 서면을 수령한 날로부터 7일 이내에 위 미지급 대금 {amt:,}원 전액을 아래 계좌로 지급하여 주실 것을 정중히 청구합니다.',after=3)
    t2=doc.add_table(rows=2,cols=4); tblb(t2)
    ws2=[Cm(3.2),Cm(5.0),Cm(4.6),Cm(CW.cm-12.8)]
    for i,(h,w) in enumerate(zip(["입금은행","계좌번호","예금주","입금액"],ws2)):
        c=t2.row_cells(0)[i]; c.width=w; shade(c); cp(c,h,align=AL.CENTER)
    dc=t2.row_cells(1)
    for i,w in enumerate(ws2): dc[i].width=w
    cp(dc[0],K["bank"],align=AL.CENTER); cp(dc[1],K["acct"],align=AL.CENTER); cp(dc[2],K["holder"],align=AL.CENTER); cp(dc[3],f"{amt:,}원",align=AL.CENTER)
    for r in t2.rows: rh(r,0.68)
    numitem("6","만약 위 기한 내에 대금이 지급되지 아니할 경우, 발신인은 부득이 민사소송 제기, 지급명령 신청 등 법적 절차에 착수할 수 밖에 없으며, 이 경우 지연손해금 및 소송비용 등이 추가로 수신인에게 부담될 수 있음을 미리 알려드립니다.")
    numitem("7","원만한 해결을 진심으로 희망하오니, 기한 내에 성실히 이행하여 주시기 바랍니다.",after=2)
    hrule(para("",before=0,after=2))
    para(datestr,AL.CENTER,size=10.5,before=date_before,after=14)
    para(K["holder"],AL.CENTER,size=11.5,bold=True,after=6+spread)
    stt=doc.add_table(rows=1,cols=3); stt.autofit=False; stt.allow_autofit=False
    set_fixed(stt,[CW.cm-8.7, 6.7, 2.0])
    _cm=OxmlElement('w:tblCellMar')        # 칸 좌우 여백 0 → 명판 그림이 잘리지 않게
    for _sd in ('top','left','bottom','right'):
        _e=OxmlElement('w:'+_sd); _e.set(qn('w:w'),'0'); _e.set(qn('w:type'),'dxa'); _cm.append(_e)
    stt._tbl.tblPr.append(_cm)
    sp=stt.cell(0,0); mc=stt.cell(0,1); sc=stt.cell(0,2)
    mc.vertical_alignment=VAL.CENTER; sc.vertical_alignment=VAL.CENTER
    mpp=mc.paragraphs[0]; mpp.alignment=AL.CENTER; mpp.add_run().add_picture(os.path.join(repo_dir,KWON[reg]["stamp"]), width=Cm(6.0))
    spp=sc.paragraphs[0]; spp.alignment=AL.CENTER; spp.add_run().add_picture(os.path.join(repo_dir,"인감_보정_투명.png"), width=Cm(1.5))
    doc.save(out_docx); return out_docx
